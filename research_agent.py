import streamlit as st
import pandas as pd
import requests
import os
import datetime
import time
import io
import re
from dotenv import load_dotenv

# --- AUTH & AI IMPORTS ---
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from google.oauth2.credentials import Credentials

# --- FILE READERS ---
try:
    from pypdf import PdfReader
except ImportError: PdfReader = None
try:
    from docx import Document
except ImportError: Document = None

# ==========================================
# 🔐 CONFIG & AUTH SETUP
# ==========================================
load_dotenv() 

def get_secret(key):
    if key in os.environ: return os.environ[key]
    try:
        if key in st.secrets: return st.secrets[key]
    except: pass
    return None

GEMINI_API_KEY = get_secret("GEMINI_API_KEY")
CLIENT_ID = get_secret("GOOGLE_CLIENT_ID")
CLIENT_SECRET = get_secret("GOOGLE_CLIENT_SECRET")
REDIRECT_URI = get_secret("REDIRECT_URI") 

SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/drive.file'] 
CLINICAL_TRIALS_API = "https://clinicaltrials.gov/api/v2/studies"

# ==========================================
# 🔑 OAUTH 2.0 LOGIN FLOW
# ==========================================
def authorize_google():
    if not CLIENT_ID or not CLIENT_SECRET or not REDIRECT_URI:
        st.error("⚠️ Missing Google OAuth Secrets.")
        return None

    client_config = {
        "web": {
            "client_id": CLIENT_ID,
            "client_secret": CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
        }
    }
    flow = Flow.from_client_config(client_config, scopes=SCOPES, redirect_uri=REDIRECT_URI)

    if "code" in st.query_params:
        code = st.query_params["code"]
        try:
            flow.fetch_token(code=code)
            creds = flow.credentials
            st.session_state["google_creds"] = {
                "token": creds.token, "refresh_token": creds.refresh_token,
                "token_uri": creds.token_uri, "client_id": creds.client_id,
                "client_secret": creds.client_secret, "scopes": creds.scopes
            }
            st.query_params.clear()
            st.rerun()
        except Exception as e: st.error(f"Login failed: {e}")

    if "google_creds" in st.session_state:
        return Credentials(**st.session_state["google_creds"])
    return None

def get_login_url():
    if not CLIENT_ID: return "#"
    client_config = {
        "web": {
            "client_id": CLIENT_ID, "client_secret": CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
        }
    }
    flow = Flow.from_client_config(client_config, scopes=SCOPES, redirect_uri=REDIRECT_URI)
    auth_url, _ = flow.authorization_url(prompt='consent')
    return auth_url

# ==========================================
# 📂 READERS (Local & Drive)
# ==========================================
def get_text_from_upload(uploaded_file):
    try:
        if uploaded_file.name.endswith('.pdf'):
            if PdfReader is None: return "Error: Install pypdf"
            reader = PdfReader(uploaded_file)
            return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif uploaded_file.name.endswith('.docx'):
            if Document is None: return "Error: Install python-docx"
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return uploaded_file.read().decode("utf-8")
    except Exception as e: return f"Error: {e}"

def extract_id_from_url(url):
    match = re.search(r'/d/([a-zA-Z0-9-_]+)', url)
    return match.group(1) if match else None

def get_text_from_drive_url(service, url):
    file_id = extract_id_from_url(url)
    if not file_id: return "Error: Could not find File ID in URL."
    
    try:
        file_meta = service.files().get(fileId=file_id, fields='mimeType, name').execute()
        mime_type = file_meta.get('mimeType')
        
        if mime_type == 'application/vnd.google-apps.document':
            content = service.files().export_media(fileId=file_id, mimeType='text/plain').execute()
            return content.decode('utf-8')
        else:
            request = service.files().get_media(fileId=file_id)
            file_content = io.BytesIO(request.execute())
            if 'pdf' in mime_type:
                if PdfReader is None: return "Error: Install pypdf"
                reader = PdfReader(file_content)
                return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
            elif 'wordprocessingml' in mime_type:
                if Document is None: return "Error: Install python-docx"
                doc = Document(file_content)
                return "\n".join([p.text for p in doc.paragraphs])
            else:
                return f"Error: Unsupported file type ({mime_type}). Only Google Docs, PDFs, and Word Docs supported."
    except Exception as e:
        return f"Error reading Drive file: {e}"

# ==========================================
# 🩺 SEARCH & ANALYSIS (FULL COLUMNS RESTORED)
# ==========================================
def fetch_recent_trials(disease_query, months_back=6):
    today = datetime.date.today()
    start_date_threshold = today - datetime.timedelta(days=30*months_back)
    params = {
        "query.term": disease_query,       
        "filter.overallStatus": "RECRUITING,NOT_YET_RECRUITING,ACTIVE_NOT_RECRUITING",
        "pageSize": 300, "sort": "StudyFirstPostDate:desc"  
    }
    try:
        headers = {"User-Agent": "ResearchAgent/OAuth"}
        response = requests.get(CLINICAL_TRIALS_API, params=params, headers=headers)
        if response.status_code != 200: return []
        data = response.json()
    except: return []
    
    studies = []
    if 'studies' in data:
        for study in data['studies']:
            # Module Extraction
            protocol = study.get('protocolSection', {})
            id_mod = protocol.get('identificationModule', {})
            stat_mod = protocol.get('statusModule', {})
            cond_mod = protocol.get('conditionsModule', {})
            design_mod = protocol.get('designModule', {})
            int_mod = protocol.get('armsInterventionsModule', {})
            loc_mod = protocol.get('contactsLocationsModule', {})
            
            # Date Filter
            start_date_str = stat_mod.get('startDateStruct', {}).get('date')
            if not start_date_str: continue

            try:
                dt = datetime.datetime.strptime(start_date_str, "%Y-%m-%d" if len(start_date_str) > 7 else "%Y-%m").date()
                if dt < start_date_threshold: continue
            except: continue

            # Disease Match
            matches = [c for c in cond_mod.get('conditions', []) if disease_query.lower() in c.lower()]
            if not matches: continue 

            # --- RESTORED FIELDS ---
            # Interventions
            interventions = int_mod.get('interventions', [])
            int_str = ", ".join([i.get('name', '') for i in interventions]) if interventions else "N/A"

            # Phase
            phases = design_mod.get('phases', [])
            phase_str = ", ".join(phases) if phases else "Not Specified"

            # Enrollment
            enrollment = design_mod.get('enrollmentInfo', {}).get('count', 'N/A')

            # Locations (Countries only to keep it short)
            locations = loc_mod.get('locations', [])
            countries = list(set([l.get('country', '') for l in locations if l.get('country')]))
            loc_str = ", ".join(countries[:5]) if countries else "N/A"

            # Completion Date
            comp_date = stat_mod.get('completionDateStruct', {}).get('date', 'N/A')

            studies.append({
                "NCT Number": id_mod.get('nctId'),
                "Study title": id_mod.get('briefTitle'),
                "Status": stat_mod.get('overallStatus'),
                "Conditions": ", ".join(matches),
                "Intervention": int_str,
                "Phase": phase_str,
                "Enrollment": enrollment,
                "Study Type": design_mod.get('studyType', 'N/A'),
                "Start Date": start_date_str,
                "Completion Date": comp_date,
                "Locations": loc_str,
                "BriefSummary": protocol.get('descriptionModule', {}).get('briefSummary', '') # For AI only
            })
    return studies

def analyze_trials(research_text, trials, api_key):
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", google_api_key=api_key, temperature=0.1)
    
    with st.status("Reading your document..."):
        summary_chain = ChatPromptTemplate.from_template("Summarize goals in 3 sentences: {text}") | llm
        research_summary = summary_chain.invoke({"text": research_text[:25000]}).content
        st.write(f"**Context:** {research_summary}")

    analyzed_data = []
    BATCH_SIZE = 20
    progress = st.progress(0)
    
    analysis_prompt = ChatPromptTemplate.from_template("""
    MY RESEARCH: {research_summary}
    TRIALS: {trials_text}
    OUTPUT: Trial_ID|Relevance(Yes/No)|Reason|Update Potential
    """)

    for i in range(0, len(trials), BATCH_SIZE):
        batch = trials[i:i+BATCH_SIZE]
        batch_txt = "\n".join([f"ID: {t['NCT Number']}\nTitle: {t['Study title']}\nSum: {t['BriefSummary'][:500]}\n---" for t in batch])
        try:
            res = (analysis_prompt | llm).invoke({"research_summary": research_summary, "trials_text": batch_txt})
            res_map = {l.split('|')[0].strip(): l.split('|')[1:] for l in res.content.split('\n') if '|' in l}
            
            for t in batch:
                if t['NCT Number'] in res_map and len(res_map[t['NCT Number']]) >= 3:
                    vals = res_map[t['NCT Number']]
                    t['AI Relevance'], t['AI Reason'], t['AI Potential'] = vals[0].strip(), vals[1].strip(), vals[2].strip()
                else: t['AI Relevance'] = "No"
                
                # CLEANUP: Remove Summary from final output to keep it clean
                if 'BriefSummary' in t: del t['BriefSummary']
                analyzed_data.append(t)
        except: pass
        progress.progress(min((i+BATCH_SIZE)/len(trials), 1.0))

    return pd.DataFrame(analyzed_data)

def upload_to_drive(service, df, filename):
    try:
        csv_buffer = io.BytesIO()
        df.to_csv(csv_buffer, index=False)
        csv_buffer.seek(0)
        file_meta = {'name': filename, 'mimeType': 'text/csv'}
        media = MediaIoBaseUpload(csv_buffer, mimetype='text/csv')
        file = service.files().create(body=file_meta, media_body=media, fields='id, webViewLink').execute()
        return file.get('webViewLink')
    except Exception as e:
        st.error(f"Upload Error: {e}")
        return None

# ==========================================
# 🚀 MAIN APP INTERFACE
# ==========================================
def main():
    st.set_page_config(page_title="Rajah", layout="wide")
    st.title("🧬 Rajah - Research AI")
    
    # Session State Init
    if 'analysis_results' not in st.session_state:
        st.session_state['analysis_results'] = None

    creds = authorize_google()
    service = build('drive', 'v3', credentials=creds) if creds else None

    # --- SIDEBAR ---
    with st.sidebar:
        st.header("1. Input Method")
        input_method = st.radio("Choose Source:", ["Upload File", "Google Doc URL"])
        
        research_text = None
        
        if input_method == "Upload File":
            uploaded_file = st.file_uploader("Drop file here", type=['pdf', 'docx', 'txt'])
            if uploaded_file:
                research_text = get_text_from_upload(uploaded_file)
        
        elif input_method == "Google Doc URL":
            if not creds: st.warning("⚠️ Login required for Drive files.")
            doc_url = st.text_input("Paste Google Doc Link")
            if doc_url and creds:
                with st.spinner("Fetching from Drive..."):
                    research_text = get_text_from_drive_url(service, doc_url)

        st.divider()
        st.header("2. Settings")
        disease = st.text_input("Disease", "Rheumatoid Arthritis")
        months = st.slider("Months Back", 1, 12, 6)
        
        st.divider()
        if not creds:
            st.link_button("🔑 Login with Google", get_login_url())
        else:
            st.success("Logged in")
            if st.button("Logout"):
                del st.session_state["google_creds"]
                st.session_state['analysis_results'] = None
                st.rerun()

    # --- MAIN CONTENT ---
    if st.button("🚀 Start Agent Analysis", type="primary"):
        st.session_state['analysis_results'] = None # Clear old results
        
        if not research_text or "Error" in research_text:
            st.error(f"Please provide a valid file. ({research_text if research_text else 'Empty'})")
        elif not GEMINI_API_KEY:
            st.error("Missing Gemini API Key!")
        else:
            trials = fetch_recent_trials(disease, months)
            if trials:
                st.info(f"Found {len(trials)} trials. Analyzing...")
                df = analyze_trials(research_text, trials, GEMINI_API_KEY)
                
                if not df.empty:
                    # Filter for High Relevance
                    relevant = df[df['AI Relevance'].str.contains("Yes", case=False, na=False)]
                    
                    # Reorder Columns for better readability
                    desired_order = [
                        'AI Relevance', 'AI Reason', 'AI Potential',
                        'NCT Number', 'Study title', 'Status', 'Phase',
                        'Intervention', 'Conditions', 'Start Date', 'Completion Date',
                        'Enrollment', 'Study Type', 'Locations'
                    ]
                    # Only select columns that actually exist
                    final_cols = [c for c in desired_order if c in relevant.columns]
                    st.session_state['analysis_results'] = relevant[final_cols]
                    st.success("Analysis Complete!")
                else:
                    st.warning("No relevant trials found.")
            else:
                st.warning("No trials found matching strict criteria.")

    # --- RESULTS DISPLAY (Persistent) ---
    if st.session_state['analysis_results'] is not None:
        final_df = st.session_state['analysis_results']
        st.subheader("High Relevance Updates")
        st.dataframe(final_df, use_container_width=True)
        
        col1, col2 = st.columns(2)
        with col1:
            csv = final_df.to_csv(index=False).encode('utf-8')
            st.download_button("📥 Download CSV", csv, "updates.csv", "text/csv")
        
        with col2:
            if service:
                if st.button("☁️ Save to Google Drive"):
                    link = upload_to_drive(service, final_df, f"Updates_{disease}.csv")
                    if link: st.success(f"Saved! [View File]({link})")
            else:
                st.caption("Login to save to Drive.")

if __name__ == "__main__":
    main()
